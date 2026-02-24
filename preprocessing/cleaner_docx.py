# backend/ai_analysis_app/preprocessing/cleaner_docx.py

import docx
import re
from typing import List, Optional
from .schema import VulnerabilityRecord

# ==============================================================================
# 🧠 EXTRACTION STRATEGY (Markdown & Structure Preserved)
# ==============================================================================
# 1. FINDING TABLES: Specialized parsing for "F-HIGH-01" style tables.
# 2. INFO TABLES: Generic parsing for Stakeholders, Tools, Scope.
# 3. NARRATIVE: Markdown capturing for Introduction, Conclusion, Methodology.

FINDING_KEYS = {
    "Vulnerability Category": "category",
    "Vulnerability Description": "description",
    "Impact": "impact",
    "Discussion of Impact": "impact",
    "Recommendation": "recommendation",
    "CVSS Score": "cvss",
    "Affected URL": "affected_url",
    "Affected IP Address": "affected_url",
    "Affected Asset": "asset_id",
    "Supporting Evidence": "raw_evidence",
    "Proof of Concept": "raw_evidence",
    "CVE": "cve_id",
    "CWE": "cve_id",
    "CVE/CWE": "cve_id",
    "Severity": "risk_level",
    "Risk": "risk_level",
    "Risk Level": "risk_level",
    "Overall Risk": "risk_level",
    "Reference": "references",
    "References": "references",
    # Handle oddly formatted tables where the Key is the severity itself
    "Critical": "risk_level",
    "High": "risk_level",
    "Medium": "risk_level",
    "Low": "risk_level",
    "Technical Severity": "risk_level",
    "Business Severity": "risk_level",
    "Likelihood": "likelihood",
    "Ease of Exploitation": "ease"
}

def clean_text(text: str) -> str:
    """Removes excess whitespace/newlines for simple fields."""
    return re.sub(r'\s+', ' ', text).strip()

def to_markdown(para) -> str:
    """Converts a docx paragraph to Markdown text based on style and content."""
    text = clean_text(para.text)
    if not text:
        return ""
    
    style = para.style.name.lower()
    
    # Headings
    if style.startswith('heading 1'): return f"# {text}"
    if style.startswith('heading 2'): return f"## {text}"
    if style.startswith('heading 3'): return f"### {text}"
    if style.startswith('heading 4'): return f"#### {text}"
    if style.startswith('heading 5'): return f"##### {text}"
    
    # Lists
    if 'list bullet' in style:
        return f"- {text}"
    if 'list number' in style or 'list paragraph' in style:
        return f"1. {text}"
        
    # Blockquotes (heuristic: "quote" or "intense quote" styles)
    if 'quote' in style:
        return f"> {text}"

    return text

def get_cell_text(cell) -> str:
    """Extracts text from a cell, preserving Markdown for internal lists/paragraphs."""
    lines = []
    for p in cell.paragraphs:
        md = to_markdown(p)
        if md:
            lines.append(md)
    return "\n\n".join(lines).strip()

# --- TABLE PARSER: SPECIALIZED FINDINGS ---
def parse_finding_table(table, filename: str) -> List[VulnerabilityRecord]:
    rows = table.rows
    if not rows: return []
    
    extracted_records = []
    current_rec = None
    
    def finalize_record(rec):
        if rec:
             # Consolidate temp description
             rec.description = "\n\n".join(rec._temp_description)
             
             # Fallback severity
             if rec.severity == 0.0 and (not rec.risk_level_str or rec.risk_level_str == "Info"):
                 if "CRITICAL" in rec.vuln_id: 
                     rec.severity = 9.5
                     rec.risk_level_str = "Critical"
                 elif "HIGH" in rec.vuln_id: 
                     rec.severity = 8.0
                     rec.risk_level_str = "High"
                 elif "MEDIUM" in rec.vuln_id: 
                     rec.severity = 5.0
                     rec.risk_level_str = "Medium"
                 elif "LOW" in rec.vuln_id: 
                     rec.severity = 2.0
                     rec.risk_level_str = "Low"
             
             # Extract OWASP link if present in references
             if rec.references and "owasp" in rec.references.lower():
                 lines = rec.references.split('\n')
                 for line in lines:
                     if "owasp.org" in line:
                         rec.owasp_link = line.strip()
                         break

             extracted_records.append(rec)

    # Helper to start a new record from a title row
    def start_new_record(title_text):
        new_rec = VulnerabilityRecord()
        new_rec.source_file = filename
        match = re.search(r'(F-[A-Z]+-\d+)\s*[:|-]\s*(.*)', title_text)
        if match:
            new_rec.vuln_id = match.group(1)
            new_rec.vuln_name = match.group(2)
        else:
            new_rec.vuln_name = title_text
        new_rec._temp_description = [] # Temporary storage
        new_rec._expecting_value_key = None # State for single-cell parsing
        return new_rec

    # Iterate rows
    for row in rows:
        first_cell_text = clean_text(row.cells[0].text)
        
        if re.search(r'^F-(CRITICAL|HIGH|MEDIUM|LOW|INFO)-\d+', first_cell_text):
            if current_rec:
                finalize_record(current_rec)
            current_rec = start_new_record(first_cell_text)
            continue
        
        if not current_rec:
             continue

        # --- Content Parsing ---
        cells = row.cells
        row_text_list = [clean_text(c.text) for c in cells]
        
        # 1. GRID HEADER DETECTION
        # Check if this row acts as a header for the next row (e.g., Overall Risk | Impact | Likelihood)
        grid_headers_map = {}
        is_header_row = False
        
        # Extended keys for grid matching
        GRID_KEYS = {
            "overall risk": "risk_level",
            "risk": "risk_level",
            "likelihood": "likelihood",
            "impact": "impact",
            "ease of exploitation": "ease",
            "cvss score": "cvss",
            "cvss": "cvss",
            "technical severity": "risk_level",
            "business severity": "risk_level",
            "cve/cwe references": "cve_id",
            "cve": "cve_id",
            "cvss vector string": "cvss_vector"
        }

        for idx, text in enumerate(row_text_list):
            t_lower = text.lower()
            if t_lower in GRID_KEYS:
                grid_headers_map[idx] = GRID_KEYS[t_lower]
                is_header_row = True
        
        if is_header_row:
            current_rec._grid_headers_for_next_row = grid_headers_map
            continue

        # 2. GRID VALUE EXTRACTION (from previous header row)
        if hasattr(current_rec, "_grid_headers_for_next_row") and current_rec._grid_headers_for_next_row:
            header_map = current_rec._grid_headers_for_next_row
            
            for idx, key in header_map.items():
                if idx < len(cells):
                    val = clean_text(cells[idx].text)
                    val_raw = get_cell_text(cells[idx])
                    
                    if key == "risk_level":
                         current_rec.risk_level_str = val
                    elif key == "likelihood":
                         current_rec.likelihood = val
                    elif key == "ease":
                         current_rec.ease = val
                    elif key == "impact":
                         # Impact often huge, append to description or specific field
                         current_rec.impact = val_raw
                    elif key == "cvss":
                         m = re.search(r'\d+(\.\d+)?', val)
                         if m:
                            if current_rec.severity == 0.0: current_rec.severity = float(m.group(0))
                            current_rec.cvss_base_score = float(m.group(0))
                    elif key == "cve_id":
                         ids = re.findall(r'(CVE-\d{4}-\d+|CWE-\d+)', val)
                         current_rec.cve_id.extend(ids)
                    elif key == "cvss_vector":
                         current_rec._temp_description.append(f"**CVSS Vector**: {val}")

            current_rec._grid_headers_for_next_row = None # Reset
            continue

        # 3. Check for merged cells (single column)
        if len(set(cells)) < 2:
            text = get_cell_text(cells[0])
            if not text: continue
            
            # Check if we were strict-waiting for a value from previous row
            if current_rec._expecting_value_key:
                v = current_rec._expecting_value_key
                value_clean = clean_text(text)
                value_raw = text
                
                if v == "affected_url":
                    current_rec.affected_url = value_clean
                elif v == "references":
                    current_rec.references = value_raw
                    current_rec._temp_description.append(f"**References**:\n{value_raw}")
                elif v == "impact":
                    current_rec.impact = value_raw
                    current_rec._temp_description.append(f"**Impact**:\n{value_raw}")
                elif v == "raw_evidence":
                    current_rec.raw_evidence = value_raw
                    current_rec._temp_description.append(f"**Evidence**:\n{value_raw}")
                elif v == "recommendation":
                    current_rec.solution = value_raw
                    current_rec._temp_description.append(f"**Recommendation**:\n{value_raw}")
                
                # Reset
                current_rec._expecting_value_key = None
                continue

            # Check if this single cell IS a Key
            matched_key = None
            text_clean_lower = text.lower().replace(":", "").strip()
            
            for k, v in FINDING_KEYS.items():
                k_lower = k.lower()
                # Use 'in' or 'startswith' for robustness (e.g. Affected URL(s) vs Affected URL)
                # But ensure it's not a narrative sentence (length check or startswith)
                if len(text_clean_lower) < 40 and k_lower in text_clean_lower:
                     matched_key = v
                     break
            
            if matched_key:
                current_rec._expecting_value_key = matched_key
                continue

            # Otherwise, it's just description text
            if len(text) > 3:
                is_artifact = re.search(r'(confidential|copyright|page \d)', text.lower())
                if not is_artifact:
                    current_rec._temp_description.append(text)
            continue
            
        # Standard Key-Value Row
        current_rec._expecting_value_key = None # Reset if we hit a normal row
        
        key_raw = clean_text(cells[0].text)
        value_raw = get_cell_text(cells[1])
        value_clean = clean_text(cells[1].text)
        
        # DEBUG
        # print(f"  [DEBUG] Row: '{key_raw}' -> '{value_clean}'")
        
        matched = False
        for k, v in FINDING_KEYS.items():
            if k.lower() in key_raw.lower():
                # If value is significantly empty, assume it's in the next row
                if len(value_clean) < 2:
                    current_rec._expecting_value_key = v
                    matched = True
                    break

                if v == "cvss":
                    m = re.search(r'\d+(\.\d+)?', value_clean)
                    if m:
                        if current_rec.severity == 0.0: current_rec.severity = float(m.group(0))
                        current_rec.cvss_base_score = float(m.group(0))
                elif v == "cve_id":
                    ids = re.findall(r'(CVE-\d{4}-\d+|CWE-\d+)', key_raw + " " + value_clean)
                    current_rec.cve_id.extend(ids)
                    current_rec._temp_description.append(f"**{key_raw}**: {value_clean}")
                elif v == "asset_id": 
                    current_rec.asset_id = value_clean
                elif v == "affected_url":
                    current_rec.affected_url = value_clean
                elif v == "recommendation": 
                    current_rec.solution = value_raw
                    current_rec._temp_description.append(f"**Recommendation**:\n{value_raw}")
                elif v == "impact":
                    current_rec.impact = value_raw
                    current_rec._temp_description.append(f"**Impact**:\n{value_raw}")
                elif v == "likelihood":
                    current_rec.likelihood = value_clean
                elif v == "ease":
                    current_rec.ease = value_clean
                elif v == "raw_evidence":
                    current_rec.raw_evidence = value_raw
                    current_rec._temp_description.append(f"**Evidence**:\n{value_raw}")
                elif v == "references":
                    current_rec.references = value_raw
                    current_rec._temp_description.append(f"**References**:\n{value_raw}")
                elif v == "risk_level": 
                    current_rec.risk_level_str = value_clean
                elif v == "description": 
                    current_rec._temp_description.insert(0, value_raw)
                else: 
                    current_rec._temp_description.append(f"**{v.title()}**:\n{value_raw}")
                matched = True
                break
        
        if not matched and len(value_clean) > 0:
            current_rec._temp_description.append(f"**{key_raw}**:\n{value_raw}")
            
    # Finalize last record
    if current_rec:
        finalize_record(current_rec)
        
    return extracted_records

def merge_finding_continuation(rec: VulnerabilityRecord, table) -> None:
    rows = table.rows
    full_description_add = []
    
    for row in rows:
        cells = row.cells
        if not any(c.text.strip() for c in cells): continue
        
        if len(set(cells)) < 2:
            text = get_cell_text(cells[0])
            if len(text) > 10:
                full_description_add.append(text)
            continue
            
        key_raw = clean_text(cells[0].text)
        value_raw = get_cell_text(cells[1])
        value_clean = clean_text(cells[1].text)
        
        matched_key = None
        for k, v in FINDING_KEYS.items():
            if k.lower() in key_raw.lower():
                matched_key = v
                break
        
        if matched_key == "cvss":
             m = re.search(r'\d+(\.\d+)?', value_clean)
             if m and rec.severity == 0.0: rec.severity = float(m.group(0))
        elif matched_key == "recommendation":
             if rec.solution: rec.solution += "\n\n" + value_raw
             else: rec.solution = value_raw
             full_description_add.append(f"**Recommendation**:\n{value_raw}")
        elif matched_key == "description":
             full_description_add.insert(0, value_raw)
        elif matched_key == "impact":
             rec.impact += "\n\n" + value_raw
             full_description_add.append(f"**Impact**:\n{value_raw}")
        elif matched_key == "raw_evidence":
             rec.raw_evidence += "\n\n" + value_raw
             full_description_add.append(f"**Evidence**:\n{value_raw}")
        elif matched_key == "risk_level":
             if not rec.risk_level_str or rec.risk_level_str == "Info":
                 rec.risk_level_str = value_clean
        elif len(key_raw) > 2 and len(value_clean) > 2:
             full_description_add.append(f"**{key_raw}**:\n{value_raw}")
        
             
    if full_description_add:
        if rec.description:
            rec.description += "\n\n" + "\n\n".join(full_description_add)
        else:
            rec.description = "\n\n".join(full_description_add)

# --- TABLE PARSER: GENERIC INFO ---
def parse_generic_table(table, filename: str, index: int) -> List[VulnerabilityRecord]:
    rows = table.rows
    if not rows: return []
    
    headers = [clean_text(c.text) for c in rows[0].cells]
    header_str = " | ".join(headers)
    
    # Analyze table type
    content_str = header_str
    for r in rows[:5]:
        content_str += " " + " ".join([c.text for c in r.cells])

    if len(set(cells for row in rows for cells in row.cells)) <= 1:
        return []

    table_type = "Generic Table"
    if ("Name" in content_str and "Role" in content_str): table_type = "Team / Stakeholders"
    elif "Vulnerability ID" in content_str and "Vulnerability Title" in content_str: table_type = "Vulnerability Summary"
    elif "Test performed" in content_str: table_type = "Methodology & Tools"
    elif "Scope" in content_str: table_type = "Scope of Activity"
    elif "CWSS" in content_str or "CVSS" in content_str: table_type = "Risk Definition"
    elif "Client" in content_str or "Version" in content_str: table_type = "Project Information"
    
    # SPECIAL LOGIC: Split Team Table
    if table_type == "Team / Stakeholders":
        quasar_rows = []
        client_rows = []
        
        for row in rows[1:]:
            row_text = " ".join([c.text for c in row.cells])
            row_cells = [clean_text(c.text) for c in row.cells]
            
            # Skip empty rows
            if not any(row_cells): continue
            
            # Determine affiliation
            if "quasar" in row_text.lower():
                quasar_rows.append(row_cells)
            else:
                client_rows.append(row_cells)
                
        records = []
        
        # Helper to build record
        def build_team_record(team_name, data_rows):
            if not data_rows: return None
            content = [f"| {' | '.join(headers)} |"]
            content.append(f"| {'--- | ' * len(headers)}")
            for r in data_rows:
                content.append(f"| {' | '.join(r)} |")
                
            rec = VulnerabilityRecord()
            rec.source_file = filename
            rec.vuln_name = f"Info: {team_name}"
            rec.asset_id = "REPORT_METADATA"
            rec.severity = 0.0 
            rec.description = "\n".join(content)
            rec.solution = f"Table Index: {index}"
            return rec

        r1 = build_team_record("Security Team (Quasar)", quasar_rows)
        if r1: records.append(r1)
        
        r2 = build_team_record("Client Team", client_rows)
        if r2: records.append(r2)
        
        return records

    # Markdown Table Construction (Default)
    content = [f"| {' | '.join(headers)} |"]
    content.append(f"| {'--- | ' * len(headers)}")
    
    for row in rows[1:]:
        row_cells = [clean_text(c.text) for c in row.cells]
        if any(row_cells):
            # Escape pipes if needed
            row_cells = [c.replace('|', '\\|') for c in row_cells]
            content.append(f"| {' | '.join(row_cells)} |")
            
    rec = VulnerabilityRecord()
    rec.source_file = filename
    rec.vuln_name = f"Info: {table_type}"
    rec.asset_id = "REPORT_METADATA"
    rec.severity = 0.0 
    rec.description = "\n".join(content)
    rec.solution = f"Table Index: {index}" 
    
    return [rec]

# --- MAIN PIPELINE ---
def process_docx_report(docx_path: str, filename: str) -> List[VulnerabilityRecord]:
    print(f"  - [DOCX Smart] 📘 Extracting RICH content from '{filename}'...")
    try:
        doc = docx.Document(docx_path)
    except Exception as e:
        print(f"  - [DOCX Smart] 🔴 Error: {e}")
        return []

    records = []
    
    # 1. PROCESS TABLES
    last_finding_record = None

    for i, table in enumerate(doc.tables):
        rows = table.rows
        if not rows: continue
        
        first_cell = clean_text(rows[0].cells[0].text)
        
        # Check if table STARTS with a finding
        if re.search(r'^F-(CRITICAL|HIGH|MEDIUM|LOW|INFO)-\d+', first_cell):
            table_records = parse_finding_table(table, filename)
            if table_records: 
                records.extend(table_records)
                last_finding_record = table_records[-1] # The last one is the "active" one for continuation
        else:
            is_continuation = False
            if last_finding_record:
                # Heuristic: Check if this looks like continuation of the LAST finding found
                # NOTE: Only simple continuation supported (not multi-finding table continuation)
                scan_rows = rows[:3]
                row_text = " ".join([c.text for r in scan_rows for c in r.cells]).lower()
                known_keys = ["description", "impact", "recommendation", "cvss vector", "overall risk"]
                if any(k in row_text for k in known_keys):
                    merge_finding_continuation(last_finding_record, table)
                    is_continuation = True
            
            if not is_continuation:
                table_records = parse_generic_table(table, filename, i)
                if table_records:
                    records.extend(table_records)
                    if table_records[0].vuln_name.startswith("Info:"): 
                        last_finding_record = None

    # 2. PROCESS NARRATIVE TEXT (Rich Markdown)
    current_section = "Introduction"
    current_text = [] # List of markdown strings
    
    for para in doc.paragraphs:
        md = to_markdown(para)
        if not md: continue
        
        style = para.style.name.lower()
        text_clean = clean_text(para.text)
        
        # New Section Detection (Headings)
        # We split larger chunks by Heading 1 or Heading 2 to keep context local
        if 'heading' in style:
            # We treat H1 and H2 as major section breaks
            is_major = 'heading 1' in style or 'heading 2' in style
            
            # If it's a major heading OR we have accumulated a LOT of text (>2000 chars), we break.
            # But relying on headings is safer for semantic grouping.
            
            if is_major and current_text:
                rec = VulnerabilityRecord()
                rec.source_file = filename
                rec.vuln_name = f"Section: {current_section}"
                rec.asset_id = "REPORT_NARRATIVE"
                rec.severity = 0.0
                rec.description = "\n\n".join(current_text)
                records.append(rec)
                current_text = []
                current_section = text_clean

            # Add the heading itself to the NEW section's text (so it has a title)
            # OR if we just flushed, start the new buffer with this heading.
            current_text.append(md)
            if is_major:
                current_section = text_clean
        else:
            current_text.append(md)
            
    # Save last section
    if current_text:
        rec = VulnerabilityRecord()
        rec.source_file = filename
        rec.vuln_name = f"Section: {current_section}"
        rec.asset_id = "REPORT_NARRATIVE"
        rec.severity = 0.0
        rec.description = "\n\n".join(current_text)
        records.append(rec)

    # 3. FALLBACK: If no records found, dump entire text
    # 3. FALLBACK: Semantic Heading-Based Chunking
    if not records:
        print("  - [DOCX Smart] ⚠️ No structure found. Applying Semantic Header Splitting.")
        
        current_chunk_title = "General Context"
        current_chunk_text = []
        current_char_count = 0
        
        for para in doc.paragraphs:
            text = clean_text(para.text)
            if not text: continue
            
            style = para.style.name.lower()
            
            # Detect Heading (New Chunk)
            if 'heading' in style or (len(text) < 60 and text.isupper()):
                # Save previous chunk if meaningful
                if current_chunk_text:
                    rec = VulnerabilityRecord()
                    rec.source_file = filename
                    rec.vuln_name = f"Section: {current_chunk_title}"
                    rec.asset_id = "REPORT_NARRATIVE"
                    rec.description = "\n\n".join(current_chunk_text)
                    rec.severity = 0.0
                    records.append(rec)
                
                # Start new chunk
                current_chunk_title = text
                current_chunk_text = [f"# {text}"] # Keep title in text
                current_char_count = len(text)
            
            else:
                current_chunk_text.append(to_markdown(para))
                current_char_count += len(text)
                
                # Split huge chunks (prevent context overflow)
                if current_char_count > 2000:
                    rec = VulnerabilityRecord()
                    rec.source_file = filename
                    rec.vuln_name = f"Section: {current_chunk_title} (Part)"
                    rec.asset_id = "REPORT_NARRATIVE"
                    rec.description = "\n\n".join(current_chunk_text)
                    records.append(rec)
                    
                    current_chunk_text = []
                    current_char_count = 0

        # Save final chunk
        if current_chunk_text:
            rec = VulnerabilityRecord()
            rec.source_file = filename
            rec.vuln_name = f"Section: {current_chunk_title}"
            rec.asset_id = "REPORT_NARRATIVE"
            rec.description = "\n\n".join(current_chunk_text)
            records.append(rec)

    print(f"  - [DOCX Smart] ✅ Extracted {len(records)} semantic chunks.")
    return records
